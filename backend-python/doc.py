from docx import Document
from docx.shared import Inches
import os

def create_or_open_word_document(file_path):
    """Funkcja do otwierania lub tworzenia nowego dokumentu Word."""
    if os.path.exists(file_path):
        doc = Document(file_path)
    else:
        doc = Document()
    return doc
    
def save_to_word(file_path, frames, ocr_results_path, text_results):
    """Funkcja do zapisywania wyników w dokumencie Word."""
    doc = create_or_open_word_document(file_path)
    
    if frames and len(frames) > 0:
        for frame in frames:
            doc.add_paragraph().add_run("Frame Image:").bold = True
            doc.add_picture(frame, width=Inches(7))

    if ocr_results_path and os.path.exists(ocr_results_path):
        try:
            with open(ocr_results_path, 'r', encoding='utf-8') as file:
                ocr_text = file.readlines()
            
            if ocr_text:
                doc.add_paragraph().add_run("OCR Results:").bold = True
                previous_paragraph = None
                for line in ocr_text:
                    clean_text = line.strip()
                    if clean_text:
                        if previous_paragraph:
                            previous_paragraph.add_run('\n' + clean_text)
                        else:
                            previous_paragraph = doc.add_paragraph(clean_text)
                
                with open(ocr_results_path, "w", encoding="utf-8") as f:
                    pass

        except Exception as e:
            print(f"Error processing OCR results: {e}")

    if text_results and isinstance(text_results, str) and text_results.strip():
        doc.add_paragraph().add_run("Audio Transcription:").bold = True
        doc.add_paragraph(text_results)

    doc.save(file_path)