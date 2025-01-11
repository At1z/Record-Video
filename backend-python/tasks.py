# tasks.py
##  celery -A tasks worker --loglevel=info --pool=solo <- odpalenie kolejki
## .\redis-server.exe <-odpalenie redis
from celery import Celery # type: ignore
from audio import convert_webm_to_wav
from speechtotext import  convert_audio_to_text
from diarization import diarize_audio
from screens import extract_different_frames
from ocr import perform_ocr_on_frames
from docx import Document
from docx.shared import Inches
import os

app = Celery("tasks", broker="redis://localhost:6379/0")

def create_or_open_word_document(file_path):
    """Funkcja do otwierania lub tworzenia nowego dokumentu Word."""
    if os.path.exists(file_path):
        doc = Document(file_path)  # Otwórz istniejący dokument
    else:
        doc = Document()  # Utwórz nowy dokument
    return doc


def save_to_word(file_path, frames, ocr_results, text_results):
    """Funkcja do zapisywania wyników w dokumencie Word."""
    doc = create_or_open_word_document(file_path)
    
    # Dodaj nagłówki do dokumentu
    doc.add_heading('Frames:', level=1)
    for frame in frames:
        doc.add_paragraph(f'Frame: {frame}')
        doc.add_paragraph().add_run("Frame Image:").bold = True
        doc.add_picture(frame, width=Inches(3))  # Możesz dostosować rozmiar zdjęcia
    
    doc.add_heading('OCR Results:', level=1)
    for ocr_result in ocr_results:
        doc.add_paragraph(f'OCR: {ocr_result}')
    
    doc.add_heading('Audio to Text Results:', level=1)

    # Assuming 'text_results' is the string returned by 'convert_audio_to_text'
    if text_results:  # Check if 'text_results' is not empty
        doc.add_paragraph('Text from audio transcription:')
        doc.add_paragraph(text_results)

    # Zapisz dokument
    doc.save(file_path)


@app.task
def process_video(file_path):

    try:
        frames = extract_different_frames(file_path, difference_threshold=0.3) ## PATH TO ONE FRAME
        ocr_results = perform_ocr_on_frames(frames) ## EXTRATED TEXT FROM ONE FRAME
        print(ocr_results)

        word_file_path = "word.docx"
        save_to_word(word_file_path, frames, ocr_results, [])

        return len(frames)
    except Exception as e:
        raise Exception(f"Error processing video: {e}")

@app.task
def process_audio(file_path):

    try:
        if file_path.endswith(".webm"):
            wav_path = convert_webm_to_wav(file_path)
            os.remove(file_path)
            file_path = wav_path
        
        results = diarize_audio(file_path)
         
        print(f"{'Speaker':<15}{'Start Time':<15}{'End Time':<15}")
        print("-" * 30)
        
        for segment in results:
            print(f"{segment['speaker']:<15}{segment['start_time']:<15}{segment['end_time']:<15}")
        
        text_results = convert_audio_to_text(file_path, results, 0.15) ## PATH TO TXT FILE WITH TEXT FROM AUDIO
        word_file_path = "word.docx"
        save_to_word(word_file_path, [], [], text_results)

        return "git"
    except Exception as e:
        raise Exception(f"Error processing audio: {e}")

