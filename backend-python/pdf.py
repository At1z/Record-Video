import os
import aspose.words as aw # type: ignore
from senderToEmail import send_file_via_email
from docx import Document
from summarization import extract_text_from_word, send_query_to_groq

def add_summarization_to_word(docx_path, summarization_text):
    """
    Dodaj podsumowanie na końcu dokumentu Word.
    """
    try:
        doc = Document(docx_path)
        doc.add_paragraph().add_run("Summarization:").bold = True
        doc.add_paragraph(summarization_text)
        doc.save(docx_path)
        print(f"Summarization added to {docx_path}")
    except Exception as e:
        print(f"Error adding summarization to Word document: {e}")

def convert_docx_to_pdf(email, docx_path, pdf_path=None):
    """
    Convert a .docx file to a .pdf file using aspose.words
    """
    text = extract_text_from_word("uploads/word.docx")
    summarization = send_query_to_groq(text)
    add_summarization_to_word("uploads/word.docx", summarization)
    

    if not os.path.exists(docx_path):
        print(f"Error: The file '{docx_path}' does not exist.")
        return

    if pdf_path is None:
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"

    try:
        doc = aw.Document(docx_path)
        doc.save(pdf_path)
        print(f"PDF successfully created at: {pdf_path}")
    # Send the PDF via email if email provided
        if email:
            send_file_via_email(
                recipient_email=email,
                file_path=pdf_path,
                subject="Here is your PDF from video-presentation",
                body="It's contains all information and summarization at the end of it ."
            )
            
    except Exception as e:
        print(f"Error occurred while converting: {e}")